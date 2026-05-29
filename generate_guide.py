"""生成 Claude / Codex 安装配置工具使用说明 (Word + PDF)"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── 内容定义 ──────────────────────────────────────────────

TITLE = "Claude / Codex 一键安装配置助手"
SUBTITLE = "Windows 小白使用说明 - 2api"

SECTIONS = [
    {
        "heading": "一、这个工具做什么",
        "paragraphs": [
            "本工具用于在 Windows 上图形化安装和配置 Claude Code 或 Codex。用户只需要双击启动界面、选择工具、粘贴 2api 令牌，再点击配置即可。",
        ],
    },
    {
        "heading": "二、准备工作",
        "bullets": [
            "Windows 10 或更高版本",
            "能正常访问网络",
            "已经有 2api 账号和令牌",
            "启动时如果弹出管理员权限提示，请点击“是”",
        ],
    },
    {
        "heading": "三、从 GitHub 下载",
        "numbered": [
            "打开本项目的 GitHub 页面",
            "点击绿色的 Code 按钮",
            "点击 Download ZIP",
            "下载完成后右键解压，不要只单独下载某一个文件",
            "解压后进入文件夹，双击「start.bat」",
        ],
    },
    {
        "heading": "四、小白使用步骤",
        "numbered": [
            "下载并解压本项目，确保所有文件在同一个文件夹中",
            "双击「start.bat」",
            "先选择「Claude Code」或「Codex 桌面端」",
            "点击「开始安装」",
            "点击「打开令牌页面」并复制 2api 令牌",
            "把令牌粘贴到界面中",
            "模型可以留空，留空会使用服务默认模型",
            "点击「一键配置」或「配置 Codex」",
            "点击「测试连接」",
            "看到连接成功后，打开新的 PowerShell 或 CMD 使用对应命令",
        ],
    },
    {
        "heading": "五、默认接口地址",
        "table": {
            "headers": ["工具", "默认接口地址"],
            "rows": [
                ["Claude Code", "https://2api.cloud/"],
                ["Codex", "https://2api.cloud/v1"],
            ],
        },
        "paragraphs_after": [
            "Claude Code 使用 Anthropic 风格接口，Codex 使用 OpenAI 风格接口，所以 Codex 地址需要带 /v1。",
        ],
    },
    {
        "heading": "六、配置成功后怎么启动",
        "code": [
            "claude",
            "claude -p \"请只回复：连接成功\" --output-format text",
            "codex",
            "codex exec \"请只回复：连接成功\"",
        ],
    },
    {
        "heading": "七、上传 GitHub 前注意",
        "bullets": [
            "不要上传自己的 2api 令牌",
            "不要上传 %USERPROFILE%\\.claude\\settings.json",
            "不要上传 %USERPROFILE%\\.codex\\auth.json",
            "不要把令牌写进 README、截图或说明书",
        ],
    },
    {
        "heading": "八、常见问题",
        "qa": [
            ("双击没有反应？", "右键「start.bat」，选择以管理员身份运行。"),
            ("模型不可用？", "把模型框留空，使用服务默认模型，或填写 2api 市场里已经订阅的模型名。"),
            ("余额不足？", "说明接口已经连通，但账号余额不足，请检查 2api 余额或套餐。"),
            ("Claude 提示 Auth conflict？", "不要同时设置 ANTHROPIC_AUTH_TOKEN 和 ANTHROPIC_API_KEY。新版工具只写 ANTHROPIC_API_KEY。"),
        ],
    },
    {
        "heading": "九、文件清单",
        "table": {
            "headers": ["文件", "用途"],
            "rows": [
                ["start.bat", "图形界面入口，普通用户双击这个文件"],
                ["ClaudeCodeGUI.ps1", "Claude Code / Codex 图形界面主脚本"],
                ["一键安装.bat", "旧版 Claude Code 命令行安装入口"],
                ["配置API.bat", "旧版 Claude Code 命令行配置入口"],
                ["install.ps1", "Claude Code 安装脚本"],
                ["configure-api.ps1", "Claude Code 命令行配置脚本"],
                ["ClaudeCodeInstaller.iss", "Inno Setup 打包脚本"],
                ["generate_guide.py", "生成 Word/PDF 说明书"],
                ["requirements.txt", "生成说明书所需依赖"],
            ],
        },
    },
]

FOOTER = "请妥善保管自己的 2api 令牌。本工具只负责安装和写入本机配置。"


# ── Word 生成 ─────────────────────────────────────────────

def make_docx(path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "微软雅黑"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for sec in SECTIONS:
        h = doc.add_heading(sec["heading"], level=2)
        for run in h.runs:
            run.font.name = "微软雅黑"
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        for text in sec.get("paragraphs", []):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

        for item in sec.get("bullets", []):
            doc.add_paragraph(item, style="List Bullet")

        for i, item in enumerate(sec.get("numbered", []), 1):
            if item.startswith("  -"):
                # 子项，作为普通缩进文本
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(2)
                p.add_run(item.strip())
            else:
                p = doc.add_paragraph(style="List Number")
                p.add_run(item)

        tbl = sec.get("table")
        if tbl:
            headers = tbl["headers"]
            rows = tbl["rows"]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h_text in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    table.rows[i + 1].cells[j].text = val
            doc.add_paragraph()

        for text in sec.get("paragraphs_after", []):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

        code_lines = sec.get("code")
        if code_lines:
            for line in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(2)

        for q, a in sec.get("qa", []):
            p = doc.add_paragraph()
            run_q = p.add_run(f"Q：{q}")
            run_q.bold = True
            p2 = doc.add_paragraph()
            p2.add_run(f"A：{a}")
            p2.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(path)
    print(f"[OK] Word 文档已生成: {path}")


# ── PDF 生成 (reportlab) ─────────────────────────────────

def register_pdf_fonts():
    regular = "C:/Windows/Fonts/msyh.ttc"
    bold = "C:/Windows/Fonts/msyhbd.ttc"
    mono = "C:/Windows/Fonts/consola.ttf"

    if os.path.exists(regular):
        pdfmetrics.registerFont(TTFont("msyh", regular))
    if os.path.exists(bold):
        pdfmetrics.registerFont(TTFont("msyh-bold", bold))
    if os.path.exists(mono):
        pdfmetrics.registerFont(TTFont("consolas", mono))


def pdf_styles():
    register_pdf_fonts()
    base = getSampleStyleSheet()
    font = "msyh" if "msyh" in pdfmetrics.getRegisteredFontNames() else "Helvetica"
    bold = "msyh-bold" if "msyh-bold" in pdfmetrics.getRegisteredFontNames() else font
    mono = "consolas" if "consolas" in pdfmetrics.getRegisteredFontNames() else font

    return {
        "title": ParagraphStyle(
            "TitleCn", parent=base["Title"], fontName=bold, fontSize=20,
            leading=26, textColor=colors.HexColor("#1A1A2E"), alignment=1,
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "SubtitleCn", parent=base["Normal"], fontName=font, fontSize=12,
            leading=18, textColor=colors.HexColor("#666666"), alignment=1,
            spaceAfter=16,
        ),
        "heading": ParagraphStyle(
            "HeadingCn", parent=base["Heading2"], fontName=bold, fontSize=13,
            leading=18, textColor=colors.HexColor("#2C3E50"), spaceBefore=8,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "BodyCn", parent=base["Normal"], fontName=font, fontSize=10,
            leading=16, spaceAfter=5,
        ),
        "bullet": ParagraphStyle(
            "BulletCn", parent=base["Normal"], fontName=font, fontSize=10,
            leading=16, leftIndent=18, firstLineIndent=-10, spaceAfter=3,
        ),
        "code": ParagraphStyle(
            "CodeCn", parent=base["Code"], fontName=mono, fontSize=9,
            leading=13, backColor=colors.HexColor("#F5F5F5"),
            leftIndent=14, rightIndent=14, spaceBefore=3, spaceAfter=3,
        ),
        "footer": ParagraphStyle(
            "FooterCn", parent=base["Normal"], fontName=font, fontSize=8,
            leading=12, textColor=colors.HexColor("#999999"), alignment=1,
            spaceBefore=12,
        ),
    }


def add_pdf_table(story, table_data, styles):
    rows = [table_data["headers"]] + table_data["rows"]
    table = Table(rows, hAlign="CENTER", colWidths=[7 * cm, 8 * cm])
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), styles["body"].fontName),
        ("FONTNAME", (0, 0), (-1, 0), styles["heading"].fontName),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6F0FA")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#1E1E1E")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BFC7D5")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(table)
    story.append(Spacer(1, 8))


def make_pdf(path):
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
    )

    story = [
        Paragraph(TITLE, styles["title"]),
        Paragraph(SUBTITLE, styles["subtitle"]),
        Spacer(1, 6),
    ]

    for sec in SECTIONS:
        story.append(Paragraph(sec["heading"], styles["heading"]))

        for text in sec.get("paragraphs", []):
            story.append(Paragraph(text, styles["body"]))

        for item in sec.get("bullets", []):
            story.append(Paragraph(f"• {item}", styles["bullet"]))

        num = 0
        for item in sec.get("numbered", []):
            if item.startswith("  -"):
                story.append(Paragraph(item.strip(), styles["bullet"]))
            else:
                num += 1
                story.append(Paragraph(f"{num}. {item}", styles["body"]))

        if sec.get("table"):
            add_pdf_table(story, sec["table"], styles)

        for text in sec.get("paragraphs_after", []):
            story.append(Paragraph(text, styles["body"]))

        for line in sec.get("code", []):
            story.append(Preformatted(line, styles["code"]))

        for q, a in sec.get("qa", []):
            story.append(Paragraph(f"<b>Q：</b>{q}", styles["body"]))
            story.append(Paragraph(f"<b>A：</b>{a}", styles["body"]))

        story.append(Spacer(1, 6))

    story.append(Paragraph(FOOTER, styles["footer"]))
    doc.build(story)
    print(f"[OK] PDF 文档已生成: {path}")


# ── 主程序 ────────────────────────────────────────────────

if __name__ == "__main__":
    docx_path = os.path.join(OUTPUT_DIR, "Claude Code 使用说明.docx")
    pdf_path = os.path.join(OUTPUT_DIR, "Claude Code 使用说明.pdf")

    try:
        make_docx(docx_path)
    except PermissionError:
        print(f"[WARN] Word 文档被占用，已跳过: {docx_path}")

    try:
        make_pdf(pdf_path)
    except PermissionError:
        print(f"[WARN] PDF 文档被占用，已跳过: {pdf_path}")

    print("\n全部生成完毕！")

